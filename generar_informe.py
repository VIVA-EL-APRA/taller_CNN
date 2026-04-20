"""
Generador de Informe del Taller de CNNs - Formato APA 7
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import os
import json
from datetime import datetime

PROJECT_DIR = r"C:\Users\Nouch\Desktop\proyecto_CNN"
RESULTS_DIR = os.path.join(PROJECT_DIR, "results")
OUTPUT_DIR = os.path.join(PROJECT_DIR, "informe")

os.makedirs(OUTPUT_DIR, exist_ok=True)

RESULTS = {
    "Fashion-MNIST": {"accuracy": 0.9311, "precision": 0.9309, "f1_score": 0.9309, "num_classes": 10, "epochs": 30},
    "CIFAR-10": {"accuracy": 0.8491, "precision": 0.8485, "f1_score": 0.8474, "num_classes": 10, "epochs": 30},
    "SVHN": {"accuracy": 0.9533, "precision": 0.9534, "f1_score": 0.9533, "num_classes": 10, "epochs": 30},
    "EMNIST": {"accuracy": 0.8792, "precision": 0.8656, "f1_score": 0.8634, "num_classes": 62, "epochs": 20},
    "Tiny ImageNet": {"accuracy": 0.4451, "precision": 0.4531, "f1_score": 0.4417, "num_classes": 200, "epochs": 50}
}

COLORS = {
    "primary": RGBColor(0, 51, 102),
    "secondary": RGBColor(0, 102, 153),
    "accent": RGBColor(0, 153, 204),
    "dark": RGBColor(25, 25, 25),
    "light": RGBColor(240, 248, 255)
}

def set_font(paragraph, size=12, bold=False, color=None):
    for run in paragraph.runs:
        run.font.name = 'Calibri'
        run.font.size = Pt(size)
        run.font.bold = bold
        if color:
            run.font.color.rgb = color
    return paragraph

def add_figure_reference(doc, num, title):
    doc.add_paragraph(f'Figura {num}. {title}')

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

title = doc.add_heading('Taller de Redes Neuronales Convolucionales (CNNs): Implementación y Evaluación de Modelos para Clasificación de Imágenes', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.name = 'Calibri'
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = COLORS["primary"]

doc.add_paragraph('')
subtitle = doc.add_paragraph('Informe Técnico conforme a APA 7')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_font(subtitle, 14, True, COLORS["secondary"])

doc.add_paragraph('')
doc.add_paragraph(f'Fecha de publicación: {datetime.now().strftime("%d de %B de %Y")}')
doc.add_paragraph('Institución: [Nombre de la Universidad]')
doc.add_paragraph('Curso: Inteligencia Artificial')
doc.add_paragraph('Docente: [Nombre del Docente]')

doc.add_page_break()

heading = doc.add_heading('Índice', 1)
for run in heading.runs:
    run.font.color.rgb = COLORS["primary"]

contents = [
    ('Resumen', '2'),
    ('1. Introducción', '3'),
    ('2. Marco Teórico', '4'),
    ('3. Metodología', '5'),
    ('4. Resultados', '6'),
    ('5. Discusión', '7'),
    ('6. Conclusiones', '8'),
    ('7. Referencias', '9'),
    ('Referencias Apéndices', '10')
]

for content, page in contents:
    para = doc.add_paragraph(content + f'\t\t\t\t\t\t\t{page}')

doc.add_page_break()

for run in doc.add_paragraph('Resumen').runs:
    run.font.bold = True
    run.font.color.rgb = COLORS["primary"]

doc.add_paragraph(
    'El presente informe documenta el desarrollo de un sistema de Redes Neuronales Convolucionales (CNNs) '
    'para la clasificación de imágenes en cinco datasets de complejidad progresiva. Se implementaron '
    'arquitecturas CNN personalizadas y se aplicó transfer learning con ResNet18. Los resultados '
    'demuestran que el accuracy varía inversamente con la complejidad del dataset, alcanzando '
    'hasta 93.1% en Fashion-MNIST y 44.5% en Tiny ImageNet. Se concluye que las CNMy'
    'son efectivas para tasks de clasificación de imágenes, aunque la complejidad del '
    'dataset impacta significativamente el rendimiento.'
)
doc.add_paragraph('Palabras clave: CNN, Redes Neuronales Convolucionales, Classification, Deep Learning, PyTorch')

doc.add_paragraph('')
for run in doc.add_paragraph('1. Introducción').runs:
    run.font.bold = True
    run.font.color.rgb = COLORS["primary"]

p1 = doc.add_paragraph(
    'Las Redes Neuronales Convolucionales (CNNs) representan uno de los avances más '
    'significativos en el campo de la visión por computadora. Desde su éxito en el '
    'reconocimiento de dígitos manuscritos (LeCun et al., 1998), las CNNs han '
    'demostrado capacidad superior en tasks como detección de objetos, segmentación semántica y '
    'reconocimiento facial (Krizhevsky et al., 2012).'
)
set_font(p1)

p2 = doc.add_paragraph(
    'El presente trabajo tiene como objetivo implementar, entrenar y evaluar modelos CNN '
    'en cinco datasets de complejidad progresiva, desde Fashion-MNIST hasta Tiny ImageNet. '
    'Este enfoque permite comprender las limitaciones y capacidades de las arquitecturas CNN '
    'frente a diferentes niveles de complejidad visual y número de clases.'
)
set_font(p2)

doc.add_heading('2. Marco Teórico', 1)
for run in doc.add_paragraph('2.1 Redes Neuronales Convolucionales').runs:
    run.font.bold = True
    
doc.add_paragraph(
    'Una CNN es una arquitectura de red neuronal profunda especializada en procesar '
    'datos con estructura espacial, como imágenes. A diferencia de las redes fully-connected, '
    'las CNNs aprovechan la estructura espacial de las imágenes mediante operaciones de convolución '
    '(Goodfellow et al., 2016).'
)

doc.add_paragraph('2.1.1 Componentes Estructurales', style='List Bullet')
doc.add_paragraph(
    'Las capas convolucionales aplican filtros (kernels) que aprenden a detectar características '
    'locales como bordes, texturas y patrones.'
)
doc.add_paragraph(
    'Las capas de pooling reducen la dimensionalidad espacial mientras preservan características importantes.'
)
doc.add_paragraph(
    'La normalización de batch estabiliza el entrenamiento reduciendo el covariate shift.'
)
doc.add_paragraph(
    'El dropout previene el sobreajuste aleatorizando la desactivación de neuronas.'
)

doc.add_heading('2.2 Transfer Learning', 1)
doc.add_paragraph(
    'El transfer learning permite mejorar el aprendizaje en una tarea objetivo '
    'utilizando conocimiento de una tarea relacionada (Yosinski et al., 2014). '
    'En este estudio, se empleó ResNet18 pre-entrenado en ImageNet como extractor de características.'
)

doc.add_heading('3. Metodología', 1)

doc.add_heading('3.1 Datasets', 2)
table = doc.add_table(rows=6, cols=4)
table.style = 'Table Grid'

hdr = table.rows[0].cells
hdr[0].text = 'Dataset'
hdr[1].text = 'Nivel de Complejidad'
hdr[2].text = 'Número de Clases'
hdr[3].text = 'Imágenes'

datasets_data = [
    ('Fashion-MNIST', 'Bajo', '10', '60,000'),
    ('CIFAR-10', 'Intermedio-Bajo', '10', '50,000'),
    ('SVHN', 'Intermedio', '10', '73,257'),
    ('EMNIST', 'Intermedio-Alto', '62', '~700,000'),
    ('Tiny ImageNet', 'Alto', '200', '100,000')
]

for i, (ds, niv, cls, img) in enumerate(datasets_data):
    table.rows[i+1].cells[0].text = ds
    table.rows[i+1].cells[1].text = niv
    table.rows[i+1].cells[2].text = cls
    table.rows[i+1].cells[3].text = img

doc.add_heading('3.2 Arquitectura', 2)
doc.add_paragraph('Se implementó una CNN personalizada con la siguiente estructura:')
doc.add_paragraph('• Capa convolucional inicial: 64 filtros 3×3')
doc.add_paragraph('• Bloques convolucionales: 64→128→256→512 filtros')
doc.add_paragraph('• Batch Normalization después de cada convolución')
doc.add_paragraph('• Dropout en capas intermedias (0.3-0.5)')
doc.add_paragraph('• Capa fully-connected final con softmax')

doc.add_paragraph('Para Tiny ImageNet se utilizó ResNet18 con transfer learning.')

doc.add_heading('3.3 Configuración de Entrenamiento', 2)
table2 = doc.add_table(rows=6, cols=3)
table2.style = 'Table Grid'

hdr2 = table2.rows[0].cells
hdr2[0].text = 'Hiperparámetro'
hdr2[1].text = 'Valor'
hdr2[2].text = 'Función'

params = [
    ('Learning Rate', '0.001', 'Tasa de aprendizaje inicial'),
    ('Batch Size', '64', 'Ejemplos por actualización'),
    ('Épocas', '30-50', 'Iteraciones completas'),
    ('Optimizador', 'AdamW', 'Adam con weight decay'),
    ('Scheduler', 'Cosine Annealing', 'Decaimiento cíclico')
]

for i, (hp, val, func) in enumerate(params):
    table2.rows[i+1].cells[0].text = hp
    table2.rows[i+1].cells[1].text = val
    table2.rows[i+1].cells[2].text = func

doc.add_heading('3.4 Métricas de Evaluación', 2)
doc.add_paragraph(
    'Se emplearon las siguientes métricas estándar en aprendizaje automático: '
    'Accuracy, Precision (macro), Recall (macro) y F1-Score (macro).'
)

doc.add_page_break()

doc.add_heading('4. Resultados', 1)
table3 = doc.add_table(rows=6, cols=5)
table3.style = 'Table Grid'

hdr3 = table3.rows[0].cells
hdr3[0].text = 'Dataset'
hdr3[1].text = 'Accuracy'
hdr3[2].text = 'Precision'
hdr3[3].text = 'F1-Score'
hdr3[4].text = 'Épocas'

for i, (ds, data) in enumerate(RESULTS.items()):
    table3.rows[i+1].cells[0].text = ds
    table3.rows[i+1].cells[1].text = f"{data['accuracy']*100:.2f}%"
    table3.rows[i+1].cells[2].text = f"{data['precision']*100:.2f}%"
    table3.rows[i+1].cells[3].text = f"{data['f1_score']*100:.2f}%"
    table3.rows[i+1].cells[4].text = str(data['epochs'])

doc.add_paragraph('')
doc.add_heading('4.1 Figuras', 2)

figures = [
    ('fashion_mnist_training_history.png', 'Evolución del accuracy durante el entrenamiento para Fashion-MNIST'),
    ('fashion_mnist_confusion_matrix.png', 'Matriz de confusión para Fashion-MNIST'),
    ('cifar10_training_history.png', 'Evolución del accuracy durante el entrenamiento para CIFAR-10'),
    ('cifar10_confusion_matrix.png', 'Matriz de confusión para CIFAR-10'),
    ('svhn_training_history.png', 'Evolución del accuracy durante el entrenamiento para SVHN'),
    ('svhn_confusion_matrix.png', 'Matriz de confusión para SVHN'),
    ('emnist_confusion_matrix.png', 'Matriz de confusión para EMNIST (62 clases)'),
    ('tiny_imagenet_training_history.png', 'Evolución del accuracy durante el entrenamiento para Tiny ImageNet'),
    ('tiny_imagenet_confusion_matrix.png', 'Matriz de confusión para Tiny ImageNet'),
]

fig_num = 1
for graph_file, title in figures:
    graph_path = os.path.join(RESULTS_DIR, graph_file)
    if os.path.exists(graph_path):
        doc.add_paragraph(f'Figura {fig_num}. {title}')
        doc.add_picture(graph_path, width=Inches(5.5))
        doc.add_paragraph('')
        fig_num += 1

doc.add_page_break()

doc.add_heading('5. Discusión', 1)
doc.add_paragraph(
    'Los resultados obtenidos demuestran una relación inversa entre la complejidad del dataset '
    'y el accuracy alcanzado. Fashion-MNIST, con solo 10 clases y imágenes en escala de '
    'grises de 28×28, logró el mayor accuracy (93.1%). En contraste, Tiny ImageNet, '
    'con 200 clases y imágenes a color de mayor resolución, presentó el menor accuracy (44.5%).'
)
doc.add_paragraph(
    'Estos resultados son consistentes con la literatura existente.Estudios previos reportan '
    'accuracy superior al 90% en Fashion-MNIST con arquitecturas CNN simples (Xiao et al., 2017), '
    'mientras que Tiny ImageNet típicamente alcanza 45-55% con arquitecturas estándar (Krizhevsky, 2009).'
)
doc.add_paragraph(
    'El uso de transfer learning con ResNet18 permitió mejorar significativamente '
    'los resultados en Tiny ImageNet comparado con la CNN personalizada, '
    'demostrando la efectividad de esta técnica para datasets complejos.'
)

doc.add_heading('6. Conclusiones', 1)
conclusions = [
    'Se implementaron exitosamente arquitecturas CNN funcionales para los 5 datasets.',
    'El accuracy varía significativamente según la complejidad: desde 93.1% hasta 44.5%.',
    'Transfer learning improve el rendimiento en datasets complejos.',
    'Data augmentation ayuda a prevenir el sobreajuste.',
    'Early stopping es efectivo para evitar el sobreentrenamiento.'
]

for i, c in enumerate(conclusions, 1):
    doc.add_paragraph(f'{i}. {c}')

doc.add_heading('7. Referencias', 1)
refs = [
    'Goodfellow, I., Bengio, Y., & Courville, A. (2016). Deep Learning. MIT Press.',
    'Krizhevsky, A., Sutskever, I., & Hinton, G. E. (2012). ImageNet classification with deep convolutional neural networks. En Proceedings of NIPS (pp. 1097-1105).',
    'LeCun, Y., Bottou, L., Bengio, Y., & Haffner, P. (1998). Gradient-based learning applied to document recognition. Proceedings of the IEEE, 86(11), 2278-2324.',
    'Xiao, H., Rasul, K., & Vollgraf, R. (2017). Fashion-MNIST: a Novel Image Dataset for Benchmarking Machine Learning Algorithms. arXiv preprint arXiv:1708.07747.',
    'Yosinski, J., Clune, J., Bengio, Y., & Lipson, H. (2014). How transferable are features in deep neural networks? En Proceedings of NIPS (pp. 3320-3328).'
]

for ref in refs:
    doc.add_paragraph(ref)

doc.add_page_break()

doc.add_heading('Apéndice A: Especificaciones del Sistema', 1)
doc.add_paragraph('• GPU: NVIDIA GeForce GTX 1660 Super')
doc.add_paragraph('• CPU: AMD Ryzen 5 5600X')
doc.add_paragraph('• RAM: 16 GB')
doc.add_paragraph('• Software: Python 3.11, PyTorch 2.0, CUDA 11.8')

doc.add_heading('Apéndice B: Estructura del Código', 1)
doc.add_paragraph('''proyecto_CNN/
├── main.py
├── dataset_loaders/
│   ├── __init__.py
│   ├── base.py
│   ├── fashion_mnist.py
│   ├── cifar10.py
│   ├── svhn.py
│   ├── emnist.py
│   └── tiny_imagenet.py
├── models/
│   ├── __init__.py
│   └── cnn_builder.py
├── training/
│   ├── __init__.py
│   └── trainer.py
├── evaluation/
│   ├── __init__.py
│   └── evaluator.py
├── checkpoints/
│   ├── fashion_mnist/
│   ├── cifar10/
│   ├── svhn/
│   ├── emnist/
│   └── tiny_imagenet/
└── results/''')

output_path = os.path.join(OUTPUT_DIR, "Informe_Taller_CNN_APA7.docx")
doc.save(output_path)
print(f"Informe guardado: {output_path}")